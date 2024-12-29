from docx import Document

# 새 문서 생성
doc = Document()

# 제목 추가
doc.add_heading('나의 첫 번째 문서', level=1)

# 본문 텍스트 추가
doc.add_paragraph('이 문서는 python-docx 패키지를 사용하여 생성되었습니다.')
doc.add_paragraph('이 패키지를 사용하면 Microsoft Word 문서를 쉽게 만들고 수정할 수 있습니다.')

# 문서 저장
doc.save('첫번째_문서.docx')

print("문서가 성공적으로 생성되었습니다!")