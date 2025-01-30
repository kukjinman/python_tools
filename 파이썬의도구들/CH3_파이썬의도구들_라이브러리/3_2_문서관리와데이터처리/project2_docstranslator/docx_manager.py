#1 os와 Document를 가져옵니다.
import os
from docx import Document

#2 docx 파일을 읽어서 텍스트를 반환하는 함수
def read_docx_files(file_path):
    doc = Document(file_path)
    doc_text = ''
    for para in doc.paragraphs:
        doc_text += para.text + '\n'

    # print(doc_text)
    return doc_text

#3 번역된 텍스트를 docx 파일로 저장하는 함수
def write_docx_files(translated_data, output_directory, file_name):
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    doc = Document()
    doc.add_paragraph(translated_data)
    doc.save(output_directory + '/translated_' + file_name)
