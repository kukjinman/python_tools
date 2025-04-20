#1 python-docx 패키지의 docx 모듈 가져오기
from docx import Document

#2 새 문서 생성
doc = Document()

#3 제목 추가
doc.add_heading('나의 write_문서', level=1)
#4 본문 텍스트 추가
doc.add_paragraph('')
main_paragraph = doc.add_paragraph('')
main_paragraph.add_run('이름 : name_var')
main_paragraph.add_run(' / ')
main_paragraph.add_run('성별 : gender_var')

#5 문서 저장
doc.save('write_문서.docx')

print("문서가 성공적으로 생성되었습니다!")